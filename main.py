 #!/usr/bin/env python3

from functools import wraps,reduce
import json
from docx import Document
import os
from os import environ as env
from werkzeug.exceptions import HTTPException
import docx2txt

from dotenv import load_dotenv, find_dotenv
from flask import Flask,request, after_this_request
from flask import jsonify
from flask import redirect,request,flash,send_from_directory
from flask import render_template
from flask import session

from flask import url_for
from authlib.flask.client import OAuth
from six.moves.urllib.parse import urlencode
#import random
import io
try:
    from PIL import Image
except ImportError:
    import Image
import pytesseract
from pytesseract import image_to_string

import re
from inscriptis import get_text
import constants
#import requests
import boto3
#from urllib.request import urlretrieve, urlparse, urljoin, urlopen
#from bs4 import *
#import urllib.request
#from bs4 import BeautifulSoup
#import itertools
#import os,shutil
import flask
from werkzeug.utils import secure_filename
import textract
from pdf2image import convert_from_path, convert_from_bytes
from pdf2image.exceptions import (
    PDFInfoNotInstalledError,
    PDFPageCountError,
    PDFSyntaxError
)

#app = Flask(__name__)
app = Flask(__name__, static_url_path='/public', static_folder='./public')




#==> Installing successful
#==> /root/.linuxbrew/Homebrew/Library/Homebrew/vendor/portable-ruby/current/bin/ruby
#ruby 2.3.7p456 (2018-03-28 revision 63024) [x86_64-linux]
#==> Add Ruby to your PATH by running:
#PATH=/root/.linuxbrew/Homebrew/Library/Homebrew/vendor/portable-ruby/current/bin:$PATH
#Don't run this as root!

#sudo docker exec -it b7e40c1ace9d /bin/bash
#sudo docker exec -it c0966fba3055 /bin/bash

import sys
#
#test -d ~/.tesseract && eval $(~/.tesseract/tessy/bin/tesseract shellenv)
#test -d ./tesseract/tessy.tesseract && eval $(./tesseract/tessy/.tesseract shellenv)
#test -r ~/.bash_profile && echo "eval \$($(brew --prefix)/bin/brew shellenv)" >>~/.bash_profile
#echo "eval \$($(brew --prefix)/bin/brew shellenv)" >>~/.profile
#!/usr/bin/env python

"""
Python-tesseract. For more information: https://github.com/madmaze/pytesseract
"""

#import shlex
#import string
#import subprocess
#import tempfile
#from contextlib import contextmanager
#from csv import QUOTE_NONE
#from distutils.version import LooseVersion
#
#from glob import iglob
#from io import BytesIO
#from os.path import normcase, normpath, realpath
#from pkgutil import find_loader
#from threading import Timer
#
#
#numpy_installed = find_loader('numpy') is not None
#if numpy_installed:
#    from numpy import ndarray
#
#pandas_installed = find_loader('pandas') is not None
#if pandas_installed:
#    import pandas as pd
#
## CHANGE THIS IF TESSERACT IS NOT IN YOUR PATH, OR IS NAMED DIFFERENTLY
##    export PATH=$PATH:/home/linuxbrew/Cellar/tesseract/4.1.0/bin
##    export PATH="$PATH:/home/linuxbrew/Cellar/tesseract/4.1.0/bin/tesseract"
#    
#tesseract_cmd = './tesseract/4.1.0/bin/tesseract'
#RGB_MODE = 'RGB'
#SUPPORTED_FORMATS = {
#    'JPEG', 'PNG', 'PBM', 'PGM', 'PPM', 'TIFF', 'BMP', 'GIF'
#}
#
##$ export PATH=$PATH:"/home/linuxbrew/Cellar/tesseract/4.1.0/bin/tesseract"
#
##export PATH="$PATH:/home/linuxbrew/Cellar/tesseract/4.1.0/bin/tesseract"
#OSD_KEYS = {
#    'Page number': ('page_num', int),
#    'Orientation in degrees': ('orientation', int),
#    'Rotate': ('rotate', int),
#    'Orientation confidence': ('orientation_conf', float),
#    'Script': ('script', str),
#    'Script confidence': ('script_conf', float)
#}
#
#
#class Output:
#    BYTES = 'bytes'
#    DATAFRAME = 'data.frame'
#    DICT = 'dict'
#    STRING = 'string'
#
#
#class PandasNotSupported(EnvironmentError):
#    def __init__(self):
#        super(PandasNotSupported, self).__init__('Missing pandas package')
#
#
#class TesseractError(RuntimeError):
#    def __init__(self, status, message):
#        self.status = status
#        self.message = message
#        self.args = (status, message)
#
#
#class TesseractNotFoundError(EnvironmentError):
#    def __init__(self):
#        super(TesseractNotFoundError, self).__init__(
#            tesseract_cmd + " is not installed or it's not in your path"
#        )
#
#
#class TSVNotSupported(EnvironmentError):
#    def __init__(self):
#        super(TSVNotSupported, self).__init__(
#            'TSV output not supported. Tesseract >= 3.05 required'
#        )
#
#
#def kill(process, code):
#    process.kill()
#    process.returncode = code
#
#
#@contextmanager
#def timeout_manager(proc, seconds=0):
#    try:
#        if not seconds:
#            yield proc.communicate()[1]
#            return
#
#        timeout_code = -1
#        timer = Timer(seconds, kill, [proc, timeout_code])
#        timer.start()
#        try:
#            _, error_string = proc.communicate()
#            yield error_string
#        finally:
#            timer.cancel()
#            if proc.returncode is timeout_code and not error_string:
#                raise RuntimeError('Tesseract process timeout')
#    finally:
#        proc.stdin.close()
#        proc.stdout.close()
#        proc.stderr.close()
#
#
#def run_once(func):
#    @wraps(func)
#    def wrapper(*args, **kwargs):
#        if wrapper._result is wrapper:
#            wrapper._result = func(*args, **kwargs)
#        return wrapper._result
#
#    wrapper._result = wrapper
#    return wrapper
#
#
#def get_errors(error_string):
#    return u' '.join(
#        line for line in error_string.decode('utf-8').splitlines()
#    ).strip()
#
#
#def cleanup(temp_name):
#    """ Tries to remove temp files by filename wildcard path. """
#    for filename in iglob(temp_name + '*' if temp_name else temp_name):
#        try:
#            os.remove(filename)
#        except OSError:
#            pass
#
#
#def prepare(image):
#    if numpy_installed and isinstance(image, ndarray):
#        image = Image.fromarray(image)
#
#    if not isinstance(image, Image.Image):
#        raise TypeError('Unsupported image object')
#
#    extension = 'PNG' if not image.format else image.format
#    if extension not in SUPPORTED_FORMATS:
#        raise TypeError('Unsupported image format/type')
#
#    if not image.mode.startswith(RGB_MODE):
#        image = image.convert(RGB_MODE)
#
#    if 'A' in image.getbands():
#        # discard and replace the alpha channel with white background
#        background = Image.new(RGB_MODE, image.size, (255, 255, 255))
#        background.paste(image, (0, 0), image)
#        image = background
#
#    image.format = extension
#    return image, extension
#
#
#def save_image(image):
#    with tempfile.NamedTemporaryFile(prefix='tess_', delete=False) as f:
#        temp_name = f.name
#
#    if isinstance(image, str):
#        return temp_name, realpath(normpath(normcase(image)))
#
#    image, extension = prepare(image)
#    input_file_name = temp_name + os.extsep + extension
#    image.save(input_file_name, format=extension, **image.info)
#    return temp_name, input_file_name
#
#
#def subprocess_args(include_stdout=True):
#    # See https://github.com/pyinstaller/pyinstaller/wiki/Recipe-subprocess
#    # for reference and comments.
#
#    kwargs = {
#        'stdin': subprocess.PIPE,
#        'stderr': subprocess.PIPE,
#        'startupinfo': None,
#        'env': os.environ
#    }
#
#    if hasattr(subprocess, 'STARTUPINFO'):
#        kwargs['startupinfo'] = subprocess.STARTUPINFO()
#        kwargs['startupinfo'].dwFlags |= subprocess.STARTF_USESHOWWINDOW
#        kwargs['startupinfo'].wShowWindow = subprocess.SW_HIDE
#
#    if include_stdout:
#        kwargs['stdout'] = subprocess.PIPE
#
#    return kwargs
#
#
#def run_tesseract(input_filename,
#                  output_filename_base,
#                  extension,
#                  lang,
#                  config='',
#                  nice=0,
#                  timeout=0):
#    cmd_args = []
#
#    if not sys.platform.startswith('win32') and nice != 0:
#        cmd_args += ('nice', '-n', str(nice))
#
#    cmd_args += (tesseract_cmd, input_filename, output_filename_base)
#
#    if lang is not None:
#        cmd_args += ('-l', lang)
#
#    if config:
#        cmd_args += shlex.split(config)
#
#    if extension and extension not in {'box', 'osd', 'tsv'}:
#        cmd_args.append(extension)
#
#    try:
#        proc = subprocess.Popen(cmd_args, **subprocess_args())
#    except OSError:
#        raise TesseractNotFoundError()
#
#    with timeout_manager(proc, timeout) as error_string:
#        if proc.returncode:
#            raise TesseractError(proc.returncode, get_errors(error_string))
#
#
#def run_and_get_output(image,
#                       extension='',
#                       lang=None,
#                       config='',
#                       nice=0,
#                       timeout=0,
#                       return_bytes=False):
#
#    temp_name, input_filename = '', ''
#    try:
#        temp_name, input_filename = save_image(image)
#        kwargs = {
#            'input_filename': input_filename,
#            'output_filename_base': temp_name + '_out',
#            'extension': extension,
#            'lang': lang,
#            'config': config,
#            'nice': nice,
#            'timeout': timeout
#        }
#
#        run_tesseract(**kwargs)
#        filename = kwargs['output_filename_base'] + os.extsep + extension
#        with open(filename, 'rb') as output_file:
#            if return_bytes:
#                return output_file.read()
#            return output_file.read().decode('utf-8').strip()
#    finally:
#        cleanup(temp_name)
#
#
#def file_to_dict(tsv, cell_delimiter, str_col_idx):
#    result = {}
#    rows = [row.split(cell_delimiter) for row in tsv.split('\n')]
#    if not rows:
#        return result
#
#    header = rows.pop(0)
#    length = len(header)
#    if len(rows[-1]) < length:
#        # Fixes bug that occurs when last text string in TSV is null, and
#        # last row is missing a final cell in TSV file
#        rows[-1].append('')
#
#    if str_col_idx < 0:
#        str_col_idx += length
#
#    for i, head in enumerate(header):
#        result[head] = list()
#        for row in rows:
#            if len(row) <= i:
#                continue
#
#            val = row[i]
#            if row[i].isdigit() and i != str_col_idx:
#                val = int(row[i])
#            result[head].append(val)
#
#    return result
#
#
#def is_valid(val, _type):
#    if _type is int:
#        return val.isdigit()
#
#    if _type is float:
#        try:
#            float(val)
#            return True
#        except ValueError:
#            return False
#
#    return True
#
#
#def osd_to_dict(osd):
#    return {
#        OSD_KEYS[kv[0]][0]: OSD_KEYS[kv[0]][1](kv[1]) for kv in (
#            line.split(': ') for line in osd.split('\n')
#        ) if len(kv) == 2 and is_valid(kv[1], OSD_KEYS[kv[0]][1])
#    }
#
#
#@run_once
#def get_tesseract_version():
#    """
#    Returns LooseVersion object of the Tesseract version
#    """
#    try:
#        return LooseVersion(
#            subprocess.check_output(
#                [tesseract_cmd, '--version'], stderr=subprocess.STDOUT
#            ).decode('utf-8').split()[1].lstrip(string.printable[10:])
#        )
#    except OSError:
#        raise TesseractNotFoundError()
#
#
#def image_to_string(image,
#                    lang=None,
#                    config='',
#                    nice=0,
#                    output_type=Output.STRING,
#                    timeout=0):
#    """
#    Returns the result of a Tesseract OCR run on the provided image to string
#    """
#    args = [image, 'txt', lang, config, nice, timeout]
#
#    return {
#        Output.BYTES: lambda: run_and_get_output(*(args + [True])),
#        Output.DICT: lambda: {'text': run_and_get_output(*args)},
#        Output.STRING: lambda: run_and_get_output(*args),
#    }[output_type]()
#
#
#def image_to_pdf_or_hocr(image,
#                         lang=None,
#                         config='',
#                         nice=0,
#                         extension='pdf',
#                         timeout=0):
#    """
#    Returns the result of a Tesseract OCR run on the provided image to pdf/hocr
#    """
#
#    if extension not in {'pdf', 'hocr'}:
#        raise ValueError('Unsupported extension: {}'.format(extension))
#    args = [image, extension, lang, config, nice, timeout, True]
#
#    return run_and_get_output(*args)
#
#
#def image_to_boxes(image,
#                   lang=None,
#                   config='',
#                   nice=0,
#                   output_type=Output.STRING,
#                   timeout=0):
#    """
#    Returns string containing recognized characters and their box boundaries
#    """
#    config += ' batch.nochop makebox'
#    args = [image, 'box', lang, config, nice, timeout]
#
#    return {
#        Output.BYTES: lambda: run_and_get_output(*(args + [True])),
#        Output.DICT: lambda: file_to_dict(
#            'char left bottom right top page\n' + run_and_get_output(*args),
#            ' ',
#            0),
#        Output.STRING: lambda: run_and_get_output(*args),
#    }[output_type]()
#
#
#def get_pandas_output(args):
#    if not pandas_installed:
#        raise PandasNotSupported()
#
#    return pd.read_csv(
#        BytesIO(run_and_get_output(*args)),
#        quoting=QUOTE_NONE,
#        sep='\t'
#    )
#
#
#def image_to_data(image,
#                  lang=None,
#                  config='',
#                  nice=0,
#                  output_type=Output.STRING,
#                  timeout=0):
#    """
#    Returns string containing box boundaries, confidences,
#    and other information. Requires Tesseract 3.05+
#    """
#
#    if get_tesseract_version() < '3.05':
#        raise TSVNotSupported()
#
#    config = '{} {}'.format('-c tessedit_create_tsv=1', config.strip()).strip()
#    args = [image, 'tsv', lang, config, nice, timeout]
#
#    return {
#        Output.BYTES: lambda: run_and_get_output(*(args + [True])),
#        Output.DATAFRAME: lambda: get_pandas_output(args + [True]),
#        Output.DICT: lambda: file_to_dict(run_and_get_output(*args), '\t', -1),
#        Output.STRING: lambda: run_and_get_output(*args),
#    }[output_type]()
#
#
#def image_to_osd(image,
#                 lang='osd',
#                 config='',
#                 nice=0,
#                 output_type=Output.STRING,
#                 timeout=0):
#    """
#    Returns string containing the orientation and script detection (OSD)
#    """
#    config = '{}-psm 0 {}'.format(
#        '' if get_tesseract_version() < '3.05' else '-',
#        config.strip()
#    ).strip()
#    args = [image, 'osd', lang, config, nice, timeout]
#
#    return {
#        Output.BYTES: lambda: run_and_get_output(*(args + [True])),
#        Output.DICT: lambda: osd_to_dict(run_and_get_output(*args)),
#        Output.STRING: lambda: run_and_get_output(*args),
#    }[output_type]()
#
#
#def main():
#    if len(sys.argv) == 2:
#        filename, lang = sys.argv[1], None
#    elif len(sys.argv) == 4 and sys.argv[1] == '-l':
#        filename, lang = sys.argv[3], sys.argv[2]
#    else:
#        sys.stderr.write('Usage: python pytesseract.py [-l lang] input_file\n')
#        exit(2)
#
#    try:
#        with Image.open(filename) as img:
#            print(image_to_string(img, lang=lang))
#    except IOError:
#        sys.stderr.write('ERROR: Could not open file "%s"\n' % filename)
#        exit(1)
#


ENV_FILE = find_dotenv()
if ENV_FILE:
    load_dotenv(ENV_FILE)

AUTH0_CALLBACK_URL = env.get(constants.AUTH0_CALLBACK_URL)
AUTH0_CLIENT_ID = env.get(constants.AUTH0_CLIENT_ID)
AUTH0_CLIENT_SECRET = env.get(constants.AUTH0_CLIENT_SECRET)
AUTH0_DOMAIN = env.get(constants.AUTH0_DOMAIN)
AUTH0_BASE_URL = AUTH0_DOMAIN
if AUTH0_BASE_URL is not None:
    AUTH0_BASE_URL = 'https://' + AUTH0_DOMAIN
AUTH0_AUDIENCE = env.get(constants.AUTH0_AUDIENCE)
if AUTH0_AUDIENCE is '':
    AUTH0_AUDIENCE = AUTH0_BASE_URL + '/userinfo'


app.secret_key = constants.SECRET_KEY
app.debug = True

@app.errorhandler(Exception)
def handle_auth_error(ex):
    response = jsonify(message=str(ex))
    response.status_code = (ex.code if isinstance(ex, HTTPException) else 500)
    return response

    
ko=AUTH0_BASE_URL
if ko is not None:
    ko=AUTH0_BASE_URL + '/oauth/token'

bac=AUTH0_BASE_URL
if bac is not None:
    bac=AUTH0_BASE_URL + '/authorize'
    

    
oauth = OAuth(app)

auth0 = oauth.register(
    'auth0',
    client_id=AUTH0_CLIENT_ID,
    client_secret=AUTH0_CLIENT_SECRET,
    api_base_url=AUTH0_BASE_URL,
    access_token_url=ko,
    authorize_url=bac,
    client_kwargs={
        'scope': 'openid profile',
    },
)


def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if constants.PROFILE_KEY not in session:
            return redirect('/login')
        return f(*args, **kwargs)

    return decorated


# Controllers API
@app.route('/')
def home():
    return render_template('home.html')



@app.route('/callback')
def callback_handling():
    auth0.authorize_access_token()
    resp = auth0.get('userinfo')
    userinfo = resp.json()

    session[constants.JWT_PAYLOAD] = userinfo
    session[constants.PROFILE_KEY] = {
        'user_id': userinfo['sub'],
        'name': userinfo['name'],
        'picture': userinfo['picture']
    }
    return redirect('/dashboard')


@app.route('/login')
def login():
    return auth0.authorize_redirect(redirect_uri=AUTH0_CALLBACK_URL, audience=AUTH0_AUDIENCE)


@app.route('/logout')
def logout():
    session.clear()
    params = {'returnTo': url_for('home', _external=True), 'client_id': AUTH0_CLIENT_ID}
    return redirect(auth0.api_base_url + '/v2/logout?' + urlencode(params))

@app.route('/dashboard')
@requires_auth
def entry():
    return render_template("autocomplete.html")

cds = os.getcwd()


Access_key_ID = env.get(constants.Access_key_ID)
Secret_access_key = env.get(constants.Secret_access_key)
print(Access_key_ID)
client = boto3.client(
    's3',
    aws_access_key_id=Access_key_ID,
    aws_secret_access_key=Secret_access_key,

)
translate = boto3.client(service_name='translate', region_name='us-east-2', use_ssl=True)

#UPLOAD_FOLDER = '/Users/berhandiclepolat/Desktop/uploads/'
UPLOAD_FOLDER = cds + "/uploads/"

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

app.secret_key = "secret key"


ALLOWED_IMAGE_EXTENSIONS = set(["tiff","txt","waw","xlsx","xls","csv","eml","epub","html", "mp3", "msg", "odt", "ogg", "jpeg", "jpg", "png", "gif", "pdf", "docx", "doc", "html","txt","pptx"])
MAX_IMAGE_FILESIZE = 16 * 1024 * 1024
ALLOWED_DOC_EXTENSIONS = set(["docx,txt"])
MAX_DOC_FILESIZE = 16 * 1024 * 1024
ALLOWED_HTML_EXTENSIONS = set(["html"])
MAX_HTML_FILESIZE = 16 * 1024 * 1024


save_path = cds + "/uploads/"


def allowed_image(imagename):
    return '.' in imagename and imagename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


@app.route('/dashboard',methods=['POST', 'GET'])
@requires_auth
def upload_file():
    if request.method == 'POST':
        image = request.files['image']
        todow = request.form.get('todow')
        sourcelang = request.form['sourcelanguage']
        targetlang = request.form['targetlanguage'] 
        
        if todow in ['OCR','PDFOCR','trans','kexkract']:
            if image and allowed_image(image.filename):
                filename = secure_filename(image.filename)
                titledoc = get_text(filename)
                titleocr = get_text(filename)
                namedok = titledoc+targetlang+".docx"
                namedok = namedok.strip()
                image.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
#                flash('file successfully uploaded')
                
                if todow in ['PDFOCR']:
                    ocrlang= request.form ["ocrlang"]
                    ocrlang = ocrlang.split(":")
     
                    ocrlang = ocrlang[1]
              
                    ocrlang = ocrlang.strip()
                    
                    conimage = convert_from_bytes(open(cds+"/uploads/"+filename, 'rb').read())
#                    conimage.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                    kocr = []
                    for i in range(len(conimage)):
                        texxt = image_to_string(Image.open(conimage[i]),lang=ocrlang)
                        kocr.append(texxt)
                        
      
        
#                        titleocr = get_text(filename)
                    titleocr = titleocr.split(".")
                    titleocr = titleocr[0]
                    titleocr = titleocr[1:]
                    titleocr = titleocr+".txt"
#                        savedocr = os.path.join(save_path, titleocr) 
#                        print(titleocr + "suiste")

        

                    imaji = io.open(titleocr,"w",encoding='utf-8')
                    imaji.write(kocr)             
                    imaji.close()
                    os.remove(cds+"/uploads/"+filename)

#                        return flask.send_file(titleocr)

                    @after_this_request
                    def lemove_ocr(response):
#                            print('After request ...')
#                        os.remove(cds + "/uploads/"+titleocr)

#                        os.remove(cds+titleocr)
                        os.remove(titleocr)
                        os.remove(conimage)
#                        os.remove(cds+filename)
                        return response

                    return flask.send_file(titleocr)
                
                
                if todow in ['OCR']:
                    ocrlang= request.form ["ocrlang"]
                    ocrlang = ocrlang.split(":")
     
                    ocrlang = ocrlang[1]
              
                    ocrlang = ocrlang.strip()
             

                    texxt = image_to_string(Image.open(image),lang=ocrlang)
      
        
#                        titleocr = get_text(filename)
                    titleocr = titleocr.split(".")
                    titleocr = titleocr[0]
                    titleocr = titleocr[1:]
                    titleocr = titleocr+".txt"
#                        savedocr = os.path.join(save_path, titleocr) 
#                        print(titleocr + "suiste")

        
                    imaji = io.open(titleocr,"w",encoding='utf-8')
                    imaji.write(texxt)             
                    imaji.close()
                    os.remove(cds+"/uploads/"+filename)

#                        return flask.send_file(titleocr)

                    @after_this_request
                    def remove_ocr(response):
#                            print('After request ...')
#                        os.remove(cds + "/uploads/"+titleocr)

#                        os.remove(cds+titleocr)
                        os.remove(titleocr)
#                        os.remove(cds+filename)
                        return response

                    return flask.send_file(titleocr)
                
                if todow in ['kexkract']:
                    titleocr = titleocr.split(".")
                    titleocr = titleocr[0]
                    titleocr = titleocr[1:]
                    titleocr = titleocr+".txt"
                    texxt = textract.process(cds+"/uploads/"+filename)
                    imaji = open(titleocr,"wb")
                    imaji.write(texxt)             
                    imaji.close()
                    os.remove(cds+"/uploads/"+filename)


                    @after_this_request
                    def remofe_ocr(response):
#                            print('After request ...')
#                        os.remove(cds + "/uploads/"+titleocr)

#                        os.remove(cds+titleocr)
                        os.remove(titleocr)
#                        os.remove(cds+filename)
                        return response

                    return flask.send_file(titleocr)
                if todow in ['trans']:

#                        doctext = doc2text.Document()
#                        doctext = doc.process()
#                        doctext = doc.extract_text()
#                        doctexttext = doc.get_text()
                    doctext = docx2txt.process(image)
#                    print(doctext)

                    
                    titledoc = titledoc.split(".")
                    titledoc = titledoc[0]
                    titledoc = titledoc[1:]                    
                    translate = boto3.client(service_name='translate', region_name='us-east-2', use_ssl=True)
                    newtext=doctext.splitlines()
                    print(newtext)
                    cleantext = [i for i in newtext if i]
                        
                    translatedsplitted = []
                    for i in range(len(cleantext)):
                        result = translate.translate_text(Text=cleantext[i], 
                                                                  SourceLanguageCode=sourcelang, TargetLanguageCode=targetlang)
                        translatedtext = result.get('TranslatedText')
                        print(translatedtext)
                        translatedsplitted.append(translatedtext)
                        
                    for k in range(len(cleantext)):
                        doctext = doctext.replace(cleantext[k],translatedsplitted[k])
                    
                    document = Document()
                    os.remove(cds+"/uploads/"+filename)

                    
                    for t in range(len(cleantext)):
                        p = document.add_paragraph(translatedsplitted[t])
                        document.save(cds+"/uploads/"+namedok)
#                        @after_this_request
#                        def remove_file(response):
#                            print('After request ...')
#                            os.remove("/Users/berhandiclepolat/Desktop/uploads/"+titledoc+targetlang+".docx")
#                            return response
                    @after_this_request
                    def remove_file(response):
#                            print('After request ...')
                        os.remove(cds+"/uploads/"+namedok)
                        return response

                    return flask.send_file(cds+"/uploads/"+namedok)

                  
                return redirect('/dashboard')
                    
        
#                    if os.path.exists(UPLOAD_FOLDER+titledoc+targetlang+".docx"):
##                        os.remove("demofile.txt")
#                
#                        os.remove(UPLOAD_FOLDER+titledoc+targetlang+".docx") 

                            
#                        return flask.send_file("/Users/berhandiclepolat/Desktop/uploads/"+titledoc+targetlang+".docx")
#                    @app.after_request 
#                    def remove_file(response): 
#                        
#                        return response 
#
#                        
#                    return render_template('autocomplete.html')
                   
                
        else:
            flash('Allowed extensions are txt, png, jpg, jpeg, gif, docx ,html')
            return redirect('/dashboard')
#            print(UPLOAD_FOLDER+titledoc+targetlang+".docx") 
            

    return render_template("autocomplete.html", todows=[{'name':'translate'},{'name':'translate'}, {'name':'OCR'}, {'name':'HTML translation'}], 
                           todow=todow,
                           targetlanguage=targetlang,ocrlang= ocrlang, sourcelanguage=sourcelang)

#@app.after_request
#def per_request_callbacks(response):
#    for func in getattr(g, 'call_after_request', ()):
#        response = func(response)
#    return response

#while True:
#    time.sleep(5)
#    shutil.rmtree("/Users/berhandiclepolat/Desktop/uploads/uploads")
#    os.mkdir("/Users/berhandiclepolat/Desktop/uploads/uploads")
#    time.sleep(1)


#43200
    
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=env.get('PORT', 3000))

                
#lsof -ti:3000 | xargs kill
    
